from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Margenes ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
AZUL_OSCURO = RGBColor(0x1F, 0x4E, 0x79)
AZUL_MEDIO  = RGBColor(0x2E, 0x75, 0xB6)
VERDE       = RGBColor(0x37, 0x86, 0x36)
ROJO        = RGBColor(0xC0, 0x00, 0x00)
GRIS        = RGBColor(0x59, 0x59, 0x59)
AMARILLO_BG = RGBColor(0xFF, 0xFF, 0xCC)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = AZUL_OSCURO
    # Linea debajo
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = AZUL_MEDIO
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text).font.size = Pt(11)
    return p

def step(num, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def tip(text, kind='info'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    colors = {'info': AZUL_MEDIO, 'ok': VERDE, 'warn': ROJO}
    icons  = {'info': 'ℹ  ', 'ok': '✔  ', 'warn': '⚠  '}
    run = p.add_run(icons.get(kind,'ℹ  ') + text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = colors.get(kind, AZUL_MEDIO)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x0A, 0x55, 0x2E)
    return p

def separador():
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("SISTEMA DE CONTROL DE INVENTARIO")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = AZUL_OSCURO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Rollos de Material")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = AZUL_MEDIO

separador()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MANUAL DE USUARIO — TABLET ANDROID")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = GRIS

separador()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Fecha: {datetime.date.today().strftime('%d/%m/%Y')}")
run.font.size = Pt(11)
run.font.color.rgb = GRIS

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. DESCRIPCION GENERAL
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. Descripcion general")
body("Esta aplicacion permite registrar rollos de material en sus ubicaciones fisicas dentro del deposito, "
     "y realizar inventarios mensuales, usando la pistola lectora de codigos de barras.")
separador()
body("La app funciona directamente desde Chrome en la tablet, sin necesidad de internet ni de instalar nada.")
tip("Los datos se guardan en la tablet automaticamente. Para pasarlos a la PC se usa la opcion Exportar.", 'info')

# ══════════════════════════════════════════════════════════════════════════════
# 2. ABRIR LA APLICACION
# ══════════════════════════════════════════════════════════════════════════════
heading1("2. Como abrir la aplicacion")
body("La primera vez, o si el acceso directo no esta disponible:")
step(1, "Abrir Chrome en la tablet.")
step(2, "En la barra de direcciones escribir la ruta del archivo:")
code("file:///storage/emulated/0/Download/inventario_android.html")
step(3, "Presionar Enter. La pantalla de inicio de sesion debe aparecer.")
separador()
body("Para crear un acceso directo en el inicio de la tablet (recomendado):")
step(1, "Con la app abierta en Chrome, tocar los tres puntos (⋮) arriba a la derecha.")
step(2, "Seleccionar  Agregar a pantalla de inicio.")
step(3, "Tocar  Agregar. Aparecera un icono en el escritorio de la tablet.")
tip("Desde el acceso directo la app abre directamente sin escribir la ruta.", 'ok')

# ══════════════════════════════════════════════════════════════════════════════
# 3. INICIO DE SESION
# ══════════════════════════════════════════════════════════════════════════════
heading1("3. Inicio de sesion")
body("Al abrir la app siempre aparece la pantalla de login. Solo el personal autorizado puede ingresar.")
separador()
step(1, "En el campo  Usuario:  escanear el codigo de usuario con la pistola  (o ingresarlo manualmente).")
step(2, "En el campo  Contraseña:  escribir la clave con el teclado de la tablet.")
step(3, "Tocar el boton  Ingresar  o presionar Enter en la contraseña.")
separador()
tip("Si los datos son incorrectos, aparece el mensaje  'Usuario o contraseña incorrectos'  en rojo.", 'warn')
tip("La sesion se mantiene activa mientras Chrome este abierto. Si se cierra Chrome, hay que volver a ingresar.", 'info')

# ══════════════════════════════════════════════════════════════════════════════
# 4. MENU PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════
heading1("4. Menu principal")
body("Luego de ingresar, aparece el menu con las siguientes opciones:")
separador()

rows = [
    ("1", "Registrar Rollo en Ubicacion", "Para asignar un rollo a una ubicacion fisica."),
    ("2", "Inventario Mensual",           "Para hacer el conteo fisico de rollos por ubicacion."),
    ("3", "Exportar Datos a PC",          "Para descargar el archivo JSON y pasarlo a la PC."),
    ("4", "Ver Datos Guardados",          "Para ver un resumen de cuantos registros hay en la tablet."),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Opcion"
hdr_cells[1].text = "Nombre"
hdr_cells[2].text = "Para que sirve"
for cell in hdr_cells:
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1F4E79')
    shd.set(qn('w:val'), 'clear')
    cell._tc.tcPr.append(shd)
for num, nombre, desc in rows:
    row_cells = table.add_row().cells
    row_cells[0].text = num
    row_cells[1].text = nombre
    row_cells[2].text = desc
    for cell in row_cells:
        cell.paragraphs[0].runs[0].font.size = Pt(10.5)

separador()
body("El operario activo se muestra en verde debajo del titulo del menu.")

# ══════════════════════════════════════════════════════════════════════════════
# 5. REGISTRAR ROLLO EN UBICACION
# ══════════════════════════════════════════════════════════════════════════════
heading1("5. Registrar Rollo en Ubicacion")
body("Se usa para asignar un rollo a una posicion concreta del deposito. "
     "Sigue tres pasos: identificar el operario, escanear la ubicacion y escanear los rollos.")

heading2("Paso 1 — Identificar operario")
step(1, "La primera vez que se usa en la sesion, la app pide el codigo de operario.")
step(2, "Escanear el badge personal con la pistola (o ingresarlo manualmente).")
step(3, "Tocar  Confirmar  o presionar Enter.")
tip("En las siguientes operaciones dentro de la misma sesion, este paso se saltea automaticamente.", 'ok')

heading2("Paso 2 — Escanear la ubicacion")
step(1, "Escanear el codigo de la ubicacion fisica (estanteria, columna, nivel).")
step(2, "Tocar  Abrir Ubicacion  o presionar Enter.")
tip("La ubicacion debe existir en el sistema. Si no existe, comunicarlo al encargado.", 'warn')

heading2("Paso 3 — Escanear los rollos")
step(1, "Escanear el codigo de barras de cada rollo uno por uno.")
step(2, "Cada escaneo agrega el rollo a la lista con la hora del registro.")
step(3, "Cuando terminaron todos los rollos de esa ubicacion, tocar  Finalizar y Guardar.")
step(4, "La app vuelve al Paso 2 para escanear la siguiente ubicacion.")
separador()
tip("Si se escanea dos veces el mismo rollo en la misma tanda, la app avisa  'DUPLICADO'  y no lo agrega.", 'warn')
tip("Si se cancela la tanda (boton Cancelar Tanda) los rollos escaneados en esa tanda NO se guardan.", 'info')

# ══════════════════════════════════════════════════════════════════════════════
# 6. INVENTARIO MENSUAL
# ══════════════════════════════════════════════════════════════════════════════
heading1("6. Inventario Mensual")
body("Se usa para el conteo fisico mensual. Registra que rollos se encontraron en cada ubicacion.")

heading2("Paso 1 — Identificar operario")
body("Igual que en Registrar Rollo. Si ya esta definido en la sesion, se salta.")

heading2("Paso 2 — Abrir el inventario")
step(1, "El nombre del inventario se completa automaticamente con el mes actual  (Ej: INV_2026-07).")
step(2, "Tocar  Abrir / Continuar  o presionar Enter.")
tip("Si el inventario ya existe, la app lo continua sin borrar lo ya registrado.", 'ok')

heading2("Paso 3 — Escanear la ubicacion")
step(1, "Escanear el codigo de la ubicacion que se esta contando.")
step(2, "Tocar  Abrir Ubicacion  o presionar Enter.")

heading2("Paso 4 — Escanear los rollos de esa ubicacion")
step(1, "Escanear cada rollo que se encuentre fisicamente en esa ubicacion.")
step(2, "Al terminar esa ubicacion, tocar  Finalizar Tanda.")
step(3, "La app vuelve al Paso 3 para seguir con la siguiente ubicacion.")
step(4, "Cuando se termina todo el deposito, tocar  Cerrar Inventario.")
separador()
tip("Si un rollo ya fue escaneado en otra ubicacion del mismo inventario, la app avisa y no lo agrega.", 'warn')

# ══════════════════════════════════════════════════════════════════════════════
# 7. EXPORTAR DATOS A LA PC
# ══════════════════════════════════════════════════════════════════════════════
heading1("7. Exportar Datos a la PC")
body("Al terminar el trabajo con la tablet, hay que pasar los datos al programa de la PC.")

heading2("En la tablet")
step(1, "Desde el menu, tocar  3 — Exportar Datos a PC.")
step(2, "Tocar  Descargar JSON.")
step(3, "Se descarga el archivo  inventario_export_FECHA.json  en la carpeta  Descargas  de la tablet.")

heading2("Pasar el archivo a la PC")
step(1, "Conectar la tablet a la PC por cable USB.")
step(2, "En la tablet, seleccionar  Transferencia de archivos  (MTP) si aparece el aviso.")
step(3, "En la PC, abrir el explorador de archivos → la tablet aparece como dispositivo externo.")
step(4, "Copiar el archivo  inventario_export_*.json  de la carpeta  Descargas  de la tablet a la carpeta del programa.")

heading2("Importar en la PC")
step(1, "Abrir el programa  inventario.py  en la PC.")
step(2, "Seleccionar  Opcion 7 — Importar de Tablet Android.")
step(3, "El programa detecta automaticamente el archivo JSON. Confirmar.")
step(4, "Los datos quedan guardados en el Excel.")
tip("Despues de importar correctamente, se puede usar el boton  Limpiar Datos Locales  en la tablet para liberar espacio.", 'info')

# ══════════════════════════════════════════════════════════════════════════════
# 8. ERRORES COMUNES
# ══════════════════════════════════════════════════════════════════════════════
heading1("8. Errores frecuentes y soluciones")
separador()

errores = [
    ("'Usuario o contraseña incorrectos'",
     "Verificar que el usuario y la contraseña esten escritos correctamente. "
     "La contraseña diferencia mayusculas y minusculas."),
    ("'DUPLICADO en tanda'",
     "El mismo rollo fue escaneado dos veces en la misma tanda. "
     "La app lo ignora automaticamente. No es necesario hacer nada."),
    ("El teclado aparece al tocar el campo de escaneo",
     "Es normal para el campo de Contraseña. En los campos de escaneo el teclado no deberia aparecer. "
     "Si aparece, tocar fuera del campo para cerrarlo y luego usar la pistola."),
    ("La app no responde al escaneo",
     "Verificar que el campo de entrada este resaltado (borde azul). "
     "Si no lo esta, tocarlo una vez y volver a escanear."),
    ("Se cerro Chrome por accidente",
     "Volver a abrir la app. Los datos ya guardados (con Finalizar) NO se pierden. "
     "Solo se pierde la tanda que estaba en curso al momento del cierre."),
]

for titulo, solucion in errores:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("Mensaje / Situacion:  ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ROJO
    run2 = p.add_run(titulo)
    run2.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after  = Pt(6)
    p2.paragraph_format.left_indent  = Cm(0.5)
    run3 = p2.add_run("Solucion:  ")
    run3.bold = True
    run3.font.size = Pt(11)
    run4 = p2.add_run(solucion)
    run4.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# 9. DATOS DE ACCESO
# ══════════════════════════════════════════════════════════════════════════════
heading1("9. Datos de acceso")
separador()

table2 = doc.add_table(rows=3, cols=2)
table2.style = 'Table Grid'
cells = table2.rows[0].cells
cells[0].text = "Campo"
cells[1].text = "Valor"
for cell in cells:
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1F4E79')
    shd.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr()
    cell._tc.tcPr.append(shd)

table2.rows[1].cells[0].text = "Usuario"
table2.rows[1].cells[1].text = "123"
table2.rows[2].cells[0].text = "Contraseña"
table2.rows[2].cells[1].text = "Operador1"

for row in table2.rows[1:]:
    for cell in row.cells:
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = AZUL_OSCURO

separador()
tip("Guardar esta informacion en un lugar seguro. No compartirla con personas no autorizadas.", 'warn')

# ── Guardar ───────────────────────────────────────────────────────────────────
doc.save("Tutorial_Inventario_Tablet.docx")
print("Archivo generado: Tutorial_Inventario_Tablet.docx")
