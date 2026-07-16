from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

AZUL    = RGBColor(0x1F, 0x4E, 0x79)
AZUL2   = RGBColor(0x2E, 0x75, 0xB6)
VERDE   = RGBColor(0x1D, 0x6A, 0x2E)
ROJO    = RGBColor(0xC0, 0x00, 0x00)
GRIS    = RGBColor(0x59, 0x59, 0x59)
NARANJA = RGBColor(0xC5, 0x5A, 0x11)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = AZUL
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1');   bot.set(qn('w:color'), '1F4E79')
    pBdr.append(bot); pPr.append(pBdr)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = AZUL2

def paso(num, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text).font.size = Pt(10.5)

def nota(text, kind='info'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.4)
    col   = {'info': AZUL2, 'ok': VERDE, 'warn': ROJO, 'tip': NARANJA}
    icon  = {'info': 'ℹ  ', 'ok': '✔  ', 'warn': '⚠  ', 'tip': '★  '}
    run = p.add_run(icon.get(kind,'ℹ  ') + text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = col.get(kind, AZUL2)

def sep():
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def tabla_pasos(filas):
    """filas = list of (icono, descripcion)"""
    t = doc.add_table(rows=len(filas), cols=2)
    t.style = 'Table Grid'
    for i, (icono, desc) in enumerate(filas):
        t.rows[i].cells[0].text = icono
        t.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(14)
        t.rows[i].cells[1].text = desc
        t.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10.5)
        t.columns[0].width = Cm(1.2)

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(50)
run = p.add_run("SISTEMA DE INVENTARIO")
run.bold = True; run.font.size = Pt(22); run.font.color.rgb = AZUL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Guia Rapida — Tablet Android")
run.bold = True; run.font.size = Pt(14); run.font.color.rgb = AZUL2

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Deposito de rollos de material  ·  {datetime.date.today().strftime('%d/%m/%Y')}")
run.font.size = Pt(10.5); run.font.color.rgb = GRIS

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. ACCESO
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Ingresar al sistema")
paso(1, "Abrir Chrome y tocar el acceso directo  INVENTARIO  en el inicio de la tablet.")
paso(2, "Ingresar el codigo de usuario con la pistola lectora.")
paso(3, "Escribir la contraseña con el teclado y tocar  Ingresar.")
nota("Si la pantalla esta en blanco, escribir en Chrome:  file:///sdcard/Download/inventario_android.html", 'info')

# ══════════════════════════════════════════════════════════════════════════════
# 2. REGISTRAR ROLLO EN UBICACION
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Registrar Rollo en Ubicacion")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.add_run("Usar esta opcion para asignar rollos a una posicion fisica del deposito.").font.size = Pt(10.5)

h2("Paso 1 — Ubicacion")
paso(1, "Desde el menu tocar  1 — Registrar Rollo en Ubicacion.")
paso(2, "Escanear el codigo de la ubicacion (debe comenzar con D: DEP1, DEP2, D077, D099, etc.).")
paso(3, "Presionar Enter o tocar  Abrir Ubicacion.")
nota("Si la ubicacion ya tiene rollos registrados, aparece la lista en verde antes de escanear.", 'tip')
nota("Desde esa lista se puede eliminar un rollo ya registrado tocando  ✕ Eliminar  y confirmando.", 'tip')

h2("Paso 2 — Escanear rollos")
paso(1, "Escanear el codigo de barras de cada rollo uno por uno.")
paso(2, "Cada rollo aparece en la lista con su hora de escaneo.")
paso(3, "Para quitar un rollo escaneado por error, tocar el boton  ✕  al lado del rollo.")
paso(4, "Al terminar con esa ubicacion tocar  Guardar y continuar con otra ubicacion.")
paso(5, "Para finalizar el trabajo del dia tocar  Guardar y Finalizar — exportar a PC.")
nota("Codigos de rollo validos: hasta 9 caracteres, NO deben comenzar con D.", 'warn')
nota("Si se escanea el mismo rollo dos veces, la app avisa  'DUPLICADO'  y no lo agrega.", 'warn')

# ══════════════════════════════════════════════════════════════════════════════
# 3. INVENTARIO MENSUAL
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Inventario Mensual")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
p.add_run("Usar esta opcion para el conteo fisico mensual de rollos por ubicacion.").font.size = Pt(10.5)

paso(1, "Desde el menu tocar  2 — Inventario Mensual.")
paso(2, "El nombre del inventario se completa solo (ej: INV_2026-07). Presionar Enter.")
paso(3, "Escanear el codigo de la ubicacion que se esta contando.")
paso(4, "Escanear cada rollo fisicamente presente en esa ubicacion.")
paso(5, "Tocar  Finalizar Tanda  y pasar a la siguiente ubicacion (volver al paso 3).")
paso(6, "Al terminar todo el deposito, tocar  Cerrar Inventario.")
nota("Si un rollo ya fue contado en otra ubicacion del mismo inventario, la app lo rechaza.", 'warn')
nota("Se puede continuar un inventario en otro momento — los datos se guardan automaticamente.", 'ok')

# ══════════════════════════════════════════════════════════════════════════════
# 4. EXPORTAR A LA PC
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Exportar datos a la PC")

h2("En la tablet")
paso(1, "Desde el menu tocar  3 — Exportar Datos a PC.")
paso(2, "Tocar  Descargar JSON.")
paso(3, "Se descarga el archivo  inventario_export_FECHA_HORA.json  en Descargas.")

h2("Pasar el archivo")
paso(1, "Conectar la tablet a la PC con el cable USB.")
paso(2, "En la tablet seleccionar  Transferencia de archivos.")
paso(3, "Copiar el archivo .json de la carpeta Descargas de la tablet a la carpeta del programa en la PC.")

h2("Importar en la PC")
paso(1, "Abrir  inventario.py  en la PC.")
paso(2, "Seleccionar opcion  7 — Importar de Tablet Android.")
paso(3, "El programa detecta el archivo automaticamente. Confirmar.")
nota("Despues de importar se puede usar  Limpiar Datos Locales  en la tablet para liberar espacio.", 'info')

# ══════════════════════════════════════════════════════════════════════════════
# 5. VALIDACIONES IMPORTANTES
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Codigos validos — referencia rapida")
sep()

t = doc.add_table(rows=4, cols=3)
t.style = 'Table Grid'

encabezados = ["Tipo de codigo", "Formato", "Ejemplos"]
for i, txt in enumerate(encabezados):
    cell = t.rows[0].cells[i]
    cell.text = txt
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10.5)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '1F4E79'); shd.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr()
    cell._tc.tcPr.append(shd)

datos = [
    ("Ubicacion",   "Empieza con D, 10 caracteres",  "DEP1C25201 / D077A10102"),
    ("Articulo",    "Hasta 9 caracteres, NO empieza con D", "FJLH786 / AB12345"),
]
for i, (tipo, fmt, ej) in enumerate(datos, start=1):
    t.rows[i].cells[0].text = tipo
    t.rows[i].cells[1].text = fmt
    t.rows[i].cells[2].text = ej
    for cell in t.rows[i].cells:
        cell.paragraphs[0].runs[0].font.size = Pt(10.5)

# fila vacia usada como nota
t.rows[3].cells[0].text = "⚠  Atencion"
t.rows[3].cells[0].paragraphs[0].runs[0].bold = True
t.rows[3].cells[0].paragraphs[0].runs[0].font.color.rgb = ROJO
t.rows[3].cells[1].merge(t.rows[3].cells[2])
t.rows[3].cells[1].text = "Si se escanea una ubicacion donde se esperaba un articulo, la app lo rechaza automaticamente."
t.rows[3].cells[1].paragraphs[0].runs[0].font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
# 6. ERRORES FRECUENTES
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Errores frecuentes")
sep()

errores = [
    ("DUPLICADO en tanda",
     "El rollo ya fue escaneado antes en esta tanda. La app lo ignora, no hace falta hacer nada."),
    ("Codigo de UBICACION escaneado",
     "Se escaneo una ubicacion (empieza con D o tiene 10 caracteres) en el campo de rollo. Escanear el codigo correcto del rollo."),
    ("Codigo invalido: debe tener 9 digitos o menos",
     "El codigo escaneado es demasiado largo. Verificar que se escaneo el rollo y no la etiqueta de la ubicacion."),
    ("La ubicacion no existe",
     "La ubicacion escaneada no esta registrada en el sistema. Comunicarlo al encargado."),
    ("Se cerro Chrome",
     "Reabrir la app. Los datos guardados con Finalizar NO se pierden. Solo se pierde la tanda en curso."),
]

for msg, sol in errores:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run("Mensaje:  ")
    r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = ROJO
    r2 = p.add_run(msg)
    r2.bold = True; r2.font.size = Pt(10.5)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(5)
    p2.paragraph_format.left_indent = Cm(0.5)
    r3 = p2.add_run("→  ")
    r3.bold = True; r3.font.size = Pt(10.5); r3.font.color.rgb = VERDE
    r4 = p2.add_run(sol)
    r4.font.size = Pt(10.5)

# ══════════════════════════════════════════════════════════════════════════════
# GUARDAR
# ══════════════════════════════════════════════════════════════════════════════
doc.save("Tutorial_Resumido_Tablet.docx")
print("Archivo generado: Tutorial_Resumido_Tablet.docx")
